from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from pathlib import Path


OUT_DIR = Path("output/ui")
OUT_DIR.mkdir(parents=True, exist_ok=True)
OUT = OUT_DIR / "Worlds_UI_and_Brand_Direction.docx"

INK = "0B0C0E"
PAPER = "F4F2ED"
WHITE = "FFFFFF"
MUTED = "686A70"
LINE = "D8D6D0"
TORCH = "FF4F2E"
TORCH_DARK = "C92D14"
MOSS = "234C3B"
SKY = "536DFE"
VIOLET = "8B5CF6"
AMBER = "F3A712"
SUCCESS = "198754"
DANGER = "D92D20"

FONT = "Arial"
DISPLAY = "Arial"
DATA = "Arial Narrow"


def rgb(hex_value):
    return RGBColor.from_string(hex_value)


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = tcPr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcPr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=140, start=160, bottom=140, end=160):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in("w:tcMar")
    if tcMar is None:
        tcMar = OxmlElement("w:tcMar")
        tcPr.append(tcMar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tcMar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tcMar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_border(cell, **kwargs):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge in ("top", "start", "bottom", "end", "insideH", "insideV"):
        if edge in kwargs:
            edge_data = kwargs.get(edge)
            tag = f"w:{edge}"
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            for key in ("val", "sz", "space", "color"):
                if key in edge_data:
                    element.set(qn(f"w:{key}"), str(edge_data[key]))


def set_repeat_table_header(row):
    trPr = row._tr.get_or_add_trPr()
    tblHeader = OxmlElement("w:tblHeader")
    tblHeader.set(qn("w:val"), "true")
    trPr.append(tblHeader)


def set_table_widths(table, widths):
    table.autofit = False
    tblPr = table._tbl.tblPr
    tblW = tblPr.find(qn("w:tblW"))
    if tblW is None:
        tblW = OxmlElement("w:tblW")
        tblPr.append(tblW)
    total = int(sum(widths) * 1440)
    tblW.set(qn("w:w"), str(total))
    tblW.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(int(width * 1440)))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = Inches(widths[i])
            tcPr = cell._tc.get_or_add_tcPr()
            tcW = tcPr.find(qn("w:tcW"))
            if tcW is None:
                tcW = OxmlElement("w:tcW")
                tcPr.append(tcW)
            tcW.set(qn("w:w"), str(int(widths[i] * 1440)))
            tcW.set(qn("w:type"), "dxa")
    tblInd = tblPr.find(qn("w:tblInd"))
    if tblInd is None:
        tblInd = OxmlElement("w:tblInd")
        tblPr.append(tblInd)
    tblInd.set(qn("w:w"), "0")
    tblInd.set(qn("w:type"), "dxa")


def set_run(run, size=10.5, color=INK, bold=False, italic=False, font=FONT, caps=False):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run.font.size = Pt(size)
    run.font.color.rgb = rgb(color)
    run.bold = bold
    run.italic = italic
    run.font.all_caps = caps
    return run


def paragraph(text="", size=10.5, color=INK, bold=False, italic=False, before=0, after=7,
              line=1.15, align=WD_ALIGN_PARAGRAPH.LEFT, font=FONT, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_together = keep
    if text:
        set_run(p.add_run(text), size=size, color=color, bold=bold, italic=italic, font=font)
    return p


def label(text, color=TORCH, after=6):
    p = paragraph(text.upper(), size=8.5, color=color, bold=True, after=after, line=1.0, keep=True)
    p.paragraph_format.keep_with_next = True
    return p


def heading(text, level=1, color=INK, before=None, after=None):
    sizes = {1: 23, 2: 15, 3: 11.5}
    defaults_before = {1: 0, 2: 13, 3: 9}
    defaults_after = {1: 9, 2: 6, 3: 4}
    p = paragraph(text, size=sizes[level], color=color, bold=True,
                  before=defaults_before[level] if before is None else before,
                  after=defaults_after[level] if after is None else after,
                  line=1.0 if level == 1 else 1.05, font=DISPLAY, keep=True)
    p.paragraph_format.keep_with_next = True
    return p


def title_block(kicker, title, deck):
    label(kicker)
    heading(title, 1)
    paragraph(deck, size=11.5, color=MUTED, after=15, line=1.25)
    rule(TORCH, 18)


def rule(color=LINE, after=10, width=16):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(after)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(width))
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def bullets(items, color=INK, compact=False):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.left_indent = Inches(0.24)
        p.paragraph_format.first_line_indent = Inches(-0.13)
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(3 if compact else 5)
        p.paragraph_format.line_spacing = 1.12
        set_run(p.add_run(item), size=10.2, color=color)


def numbered(items):
    numbering = doc.part.numbering_part.element
    abstract_ids = [int(n.get(qn("w:abstractNumId"))) for n in numbering.findall(qn("w:abstractNum"))]
    num_ids = [int(n.get(qn("w:numId"))) for n in numbering.findall(qn("w:num"))]
    abstract_id = max(abstract_ids, default=0) + 1
    num_id = max(num_ids, default=0) + 1

    abstract = OxmlElement("w:abstractNum")
    abstract.set(qn("w:abstractNumId"), str(abstract_id))
    multi = OxmlElement("w:multiLevelType")
    multi.set(qn("w:val"), "singleLevel")
    abstract.append(multi)
    lvl = OxmlElement("w:lvl")
    lvl.set(qn("w:ilvl"), "0")
    start = OxmlElement("w:start")
    start.set(qn("w:val"), "1")
    num_fmt = OxmlElement("w:numFmt")
    num_fmt.set(qn("w:val"), "decimal")
    lvl_text = OxmlElement("w:lvlText")
    lvl_text.set(qn("w:val"), "%1.")
    suff = OxmlElement("w:suff")
    suff.set(qn("w:val"), "tab")
    ppr = OxmlElement("w:pPr")
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "num")
    tab.set(qn("w:pos"), "430")
    tabs.append(tab)
    ind = OxmlElement("w:ind")
    ind.set(qn("w:left"), "430")
    ind.set(qn("w:hanging"), "250")
    ppr.extend([tabs, ind])
    lvl.extend([start, num_fmt, lvl_text, suff, ppr])
    abstract.append(lvl)
    numbering.append(abstract)
    num = OxmlElement("w:num")
    num.set(qn("w:numId"), str(num_id))
    abstract_ref = OxmlElement("w:abstractNumId")
    abstract_ref.set(qn("w:val"), str(abstract_id))
    num.append(abstract_ref)
    numbering.append(num)

    for item in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(5)
        p.paragraph_format.line_spacing = 1.15
        pPr = p._p.get_or_add_pPr()
        numPr = OxmlElement("w:numPr")
        ilvl = OxmlElement("w:ilvl")
        ilvl.set(qn("w:val"), "0")
        numId = OxmlElement("w:numId")
        numId.set(qn("w:val"), str(num_id))
        numPr.extend([ilvl, numId])
        pPr.append(numPr)
        set_run(p.add_run(item), size=10.2)


def callout(label_text, body, fill=PAPER, accent=TORCH):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [7.15])
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_margins(cell, 180, 210, 180, 210)
    set_cell_border(cell, start={"val":"single", "sz":"22", "color":accent})
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_run(p.add_run(label_text.upper()), size=8.2, color=accent, bold=True)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    p2.paragraph_format.line_spacing = 1.2
    body_color = WHITE if fill == INK else INK
    set_run(p2.add_run(body), size=11, color=body_color, bold=True)
    paragraph("", after=7)


def data_table(headers, rows, widths, header_fill=INK):
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, widths)
    border = {"val":"single", "sz":"5", "color":LINE}
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, text in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_shading(cell, header_fill)
        set_cell_margins(cell, 120, 140, 120, 140)
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(text.upper()), size=7.8, color=WHITE, bold=True)
        set_cell_border(cell, top=border, start=border, bottom=border, end=border)
    for ridx, row_data in enumerate(rows):
        cells = table.add_row().cells
        for i, text in enumerate(row_data):
            cell = cells[i]
            set_cell_shading(cell, WHITE if ridx % 2 == 0 else "FAF9F6")
            set_cell_margins(cell, 125, 140, 125, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            set_run(p.add_run(str(text)), size=8.8, color=INK, bold=(i == 0))
            set_cell_border(cell, top=border, start=border, bottom=border, end=border)
    paragraph("", after=5)
    return table


def swatch_table(swatches):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_table_widths(table, [1.08, 2.05, 1.35, 2.67])
    for i, h in enumerate(("Swatch", "Token", "Value", "Role")):
        c = table.rows[0].cells[i]
        set_cell_shading(c, INK)
        set_cell_margins(c, 100, 130, 100, 130)
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        set_run(p.add_run(h.upper()), size=7.8, color=WHITE, bold=True)
    set_repeat_table_header(table.rows[0])
    for name, value, role in swatches:
        cells = table.add_row().cells
        set_cell_shading(cells[0], value)
        cells[0].paragraphs[0].paragraph_format.space_after = Pt(18)
        for cell in cells:
            set_cell_margins(cell, 145, 140, 145, 140)
            set_cell_border(cell, bottom={"val":"single","sz":"5","color":LINE})
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        values = ("", name, f"#{value}", role)
        for i, text in enumerate(values):
            p = cells[i].paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            set_run(p.add_run(text), size=8.8, color=INK, bold=(i == 1))
    return table


def page_break():
    doc.add_page_break()


doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.68)
section.bottom_margin = Inches(0.62)
section.left_margin = Inches(0.68)
section.right_margin = Inches(0.67)
section.header_distance = Inches(0.28)
section.footer_distance = Inches(0.28)

# Default and semantic styles
normal = doc.styles["Normal"]
normal.font.name = FONT
normal._element.rPr.rFonts.set(qn("w:ascii"), FONT)
normal._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
normal.font.size = Pt(10.5)
normal.font.color.rgb = rgb(INK)
normal.paragraph_format.space_after = Pt(7)
normal.paragraph_format.line_spacing = 1.15

for style_name in ("List Bullet", "List Number"):
    st = doc.styles[style_name]
    st.font.name = FONT
    st._element.rPr.rFonts.set(qn("w:ascii"), FONT)
    st._element.rPr.rFonts.set(qn("w:hAnsi"), FONT)
    st.font.size = Pt(10.2)

# Header/footer
header = section.header
hp = header.paragraphs[0]
hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
hp.paragraph_format.space_after = Pt(0)
set_run(hp.add_run("WORLDS / PRODUCT DESIGN DIRECTION"), size=7.4, color=MUTED, bold=True)
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
fp.paragraph_format.space_after = Pt(0)
set_run(fp.add_run("RUN THE STORY.  /  JULY 2026  /  "), size=7.2, color=MUTED, bold=True)
fld = OxmlElement("w:fldSimple")
fld.set(qn("w:instr"), "PAGE")
fp._p.append(fld)

# 1 Cover
paragraph("", after=90)
label("Brand & interface direction / v1.0", color=TORCH, after=12)
paragraph("WORLDS", size=45, color=INK, bold=True, after=2, line=0.9, font=DISPLAY)
paragraph("RUN THE STORY.", size=20, color=INK, bold=True, after=18, line=1.0, font=DISPLAY)
rule(TORCH, 25, 24)
paragraph("An intentional visual system for a premium authored running game.", size=15, color=MUTED, after=12, line=1.25)
paragraph("Built for the moment before the run, the glance during it, and the memory afterward.", size=11.5, color=INK, after=90, line=1.35)
callout("North star", "Worlds should feel like a premium entertainment platform with the operational confidence of a best-in-class running app.", fill=INK, accent=TORCH)
paragraph("Prepared for product, design, engineering, story, audio and demo production.", size=8.5, color=MUTED, after=0)

# 2 Direction
page_break()
title_block("01 / Direction", "The interface gets out of the story's way.", "Worlds is cinematic in content and athletic in operation. The product must feel authored, immediate and trustworthy without borrowing the visual clutter of conventional games.")
callout("Design thesis", "Use NRC-level restraint for structure and run-time usability. Put the drama in the world art, voice performance and responsive states - not in ornamental interface chrome.")
heading("The three jobs", 2)
data_table(["Moment", "User need", "Design response"], [
    ("Before the run", "Choose quickly and feel anticipation.", "Editorial world art, one clear recommendation, plain duration and readiness information."),
    ("During the run", "Understand the situation in under one second.", "Route first, one active story objective, high-contrast metrics and voice-led feedback."),
    ("After the run", "Feel that this run mattered.", "A narrative recap tied to exact movement, decisions, consequences and character state."),
], [1.35, 2.45, 3.35])
heading("What this is not", 2)
bullets([
    "Not a fantasy-themed running dashboard.",
    "Not an audiobook player with GPS added beneath it.",
    "Not a chat interface disguised as a game director.",
    "Not an endless catalogue of interchangeable AI stories.",
    "Not a stat-heavy training tool competing with the narrative for attention.",
], compact=True)

# 3 Brand platform
page_break()
title_block("02 / Brand platform", "A world you physically enter.", "The brand promise is agency: real movement creates authored consequences. Every line of copy and every product state should reinforce that contract.")
data_table(["Element", "Approved expression"], [
    ("Category", "Audio-first running game"),
    ("Core promise", "Premium authored adventures where your run controls the story."),
    ("Tagline", "Run the Story."),
    ("Philosophy", "The world is authored. The experience is adaptive."),
    ("Creator framing", "The writer creates the world. The runner determines the story."),
    ("Proof", "Direction chooses paths. Pace changes encounters. Stops uncover moments. Failure creates story."),
], [1.55, 5.6])
heading("Brand character", 2)
data_table(["We are", "We are not"], [
    ("Cinematic", "Decorative"),
    ("Immediate", "Hyperactive"),
    ("Confident", "Aggressive"),
    ("Authored", "Synthetic"),
    ("Athletic", "Clinical"),
    ("Consequential", "Punitive"),
], [3.575, 3.575])
callout("The emotional test", "A runner should feel summoned, not onboarded; accompanied, not instructed; remembered, not scored.")

# 4 Voice
page_break()
title_block("03 / Voice", "Short. Human. Inside the world.", "Product language is calm and exact. Character language may be emotional, but the interface never performs excitement with punctuation, hype or generic fantasy phrasing.")
heading("Interface voice", 2)
data_table(["Principle", "Use", "Avoid"], [
    ("Direct", "Start run", "Begin your epic journey"),
    ("Situated", "The northern road is still open.", "A new path has been unlocked!"),
    ("Consequential", "Mara remembers you stayed.", "+10 relationship points"),
    ("Physical", "Hold this pace.", "Sprint for 20 seconds"),
    ("Specific", "You reached the gate in 04:12.", "Great job, runner!"),
], [1.25, 2.8, 3.1])
heading("Mechanics become dialogue", 2)
data_table(["System event", "Diegetic line"], [
    ("Turn left", "The forest. Good. Stay beneath the branches."),
    ("Increase pace", "They heard us. Move."),
    ("Slow down", "Torches ahead. Quiet your steps."),
    ("Stop", "Wait. There is someone under the bridge."),
    ("Fail challenge", "The gate closes. Mara takes the lower path alone."),
], [2.0, 5.15])
heading("Capitalization and punctuation", 2)
bullets(["Sentence case everywhere except the wordmark and small utility labels.", "One exclamation mark is already one too many in product UI.", "Use numerals for time, pace and distance. Write narrative quantities naturally in dialogue.", "Never expose system language such as event detected, branch selected or AI response."], compact=True)

# 5 Logo
page_break()
title_block("04 / Identity", "The name is the mark.", "Worlds needs the confidence of a media platform, not a fantasy crest. Begin with a strong wordmark and a single route symbol; add no lore-specific ornament to the parent brand.")
paragraph("WORLDS", size=42, color=INK, bold=True, after=2, line=0.9, font=DISPLAY)
paragraph("RUN THE STORY.", size=12, color=TORCH, bold=True, after=20)
heading("Wordmark direction", 2)
bullets([
    "Set WORLDS in a tightly tracked geometric grotesk, uppercase, heavy weight.",
    "Customise the W with a subtle inward route cut; it should read as typography before symbolism.",
    "Keep the silhouette horizontal and compact enough for a top bar, app launch and audio cover tile.",
    "Use black, warm white or Torch only. No gradients, bevels, outlines or genre textures.",
])
heading("Symbol direction", 2)
paragraph("A continuous line that changes direction once and terminates in a small point. It suggests route, choice and forward motion without becoming a map pin, compass or play button.", after=9)
data_table(["Lockup", "Use"], [
    ("Wordmark", "Primary: launch, campaign art, partner materials, store presence."),
    ("Route symbol", "Small spaces: app icon, map marker, audio speaking state."),
    ("Wordmark + tagline", "Brand moments only. Never inside the active-run HUD."),
], [1.75, 5.4])
callout("Guardrail", "Individual story worlds may have expressive title treatments. The Worlds parent identity remains neutral enough to host horror, fantasy, science fiction and mystery with equal authority.")

# 6 Type sheet
page_break()
title_block("05 / Typography", "Athletic scale. Editorial rhythm.", "The system uses one practical sans family and one compressed data face. Hierarchy comes from scale, weight and spacing - never from decorative type effects.")
label("Primary family / Arial")
paragraph("The forest could hide us.", size=28, color=INK, bold=True, after=4, line=1.0)
paragraph("Arial Bold / 28 pt / -1% tracking equivalent", size=8.5, color=MUTED, after=12)
paragraph("Mara remembers that you stayed.", size=17, color=INK, bold=True, after=4)
paragraph("Arial Bold / 17 pt / sentence case", size=8.5, color=MUTED, after=12)
paragraph("The road narrows beyond the gate. You have enough light to reach the ridge, but not enough to return the same way.", size=11, color=INK, after=4, line=1.3)
paragraph("Arial Regular / 11 pt / 130% line height", size=8.5, color=MUTED, after=14)
label("Performance numerals / Arial Narrow")
paragraph("04:12     5'28\"     3.84 KM", size=27, color=INK, bold=True, font=DATA, after=4)
paragraph("Arial Narrow Bold / tabular-style data treatment / never for prose", size=8.5, color=MUTED, after=14)
heading("Production recommendation", 2)
paragraph("Use Arial consistently for the hackathon interface and presentation so the typography renders identically across the working toolchain. Pair it with Arial Narrow only for live performance numerals. If the product later migrates to a bundled family such as Inter, preserve this scale, weight and spacing system rather than redesigning the hierarchy.")

# 7 Type scale
page_break()
title_block("06 / Type system", "Every size has a job.", "The scale is intentionally narrow. Screens should rarely need more than four roles at once.")
data_table(["Token", "Spec", "Use"], [
    ("display-01", "40 / 40, Heavy", "Campaign hero, completion moment"),
    ("display-02", "32 / 34, Heavy", "World title, primary route prompt"),
    ("title-01", "24 / 28, Bold", "Screen title"),
    ("title-02", "20 / 24, Demi", "Section title, character state"),
    ("body-01", "16 / 22, Medium", "Primary mobile body and prompts"),
    ("body-02", "14 / 20, Regular", "Supporting copy"),
    ("label-01", "12 / 14, Bold, +4%", "Controls and metadata"),
    ("micro-01", "11 / 13, Demi, +6%", "Map labels and captions"),
    ("metric-01", "44 / 42, DIN Bold", "Primary active metric"),
    ("metric-02", "24 / 24, DIN Bold", "Secondary metric"),
], [1.35, 1.85, 3.95])
heading("Rules", 2)
bullets([
    "Use sentence case for all titles except the wordmark and utility labels.",
    "Keep campaign titles to two lines; art direction must accommodate real title length.",
    "Use Heavy only for decisive moments. A screen filled with heavy type has no hierarchy.",
    "Never place text directly over a busy face or plot-critical artwork without a controlled scrim.",
    "Metrics use fixed-width numeral behavior to prevent layout shifting while the runner moves.",
])

# 8 Color sheet
page_break()
title_block("07 / Color sheet", "One signal in a quiet field.", "The core palette is warm rather than digital. Torch marks action, live state and narrative consequence; it is not a decorative accent sprayed across every screen.")
swatch_table([
    ("Ink", INK, "Primary background, text, control surfaces"),
    ("Paper", PAPER, "Warm light background, cards, editorial space"),
    ("White", WHITE, "High-contrast text and clean surfaces"),
    ("Torch", TORCH, "Primary action, active route, live audio state"),
    ("Ash", MUTED, "Secondary text and inactive metadata"),
    ("Line", LINE, "Rules, dividers and quiet outlines"),
])
paragraph("", after=7)
callout("Usage ratio", "Aim for 70% Ink or Paper, 20% imagery, 8% neutral structure and no more than 2% Torch on a typical product screen.")
heading("Core combinations", 2)
data_table(["Combination", "Use"], [
    ("Paper / Ink", "Library, world detail, recap reading"),
    ("Ink / White", "Active run, launch, cinematic transitions"),
    ("Torch / Ink", "Primary button and selected route"),
    ("Ink / Torch", "Live state indicator, small highlight only"),
], [2.25, 4.9])

# 9 Functional color
page_break()
title_block("08 / Functional color", "Color communicates state, not genre.", "Story artwork may vary by world. Interface colors stay stable so the runner learns what is actionable, live, successful or dangerous.")
swatch_table([
    ("Route / selected", TORCH, "Current path and primary action"),
    ("Route / alternate", SKY, "Available but unselected route"),
    ("Discovery", VIOLET, "Optional scene or unresolved mystery"),
    ("Resource", AMBER, "Supply, time or scarce opportunity"),
    ("Success", SUCCESS, "Challenge resolved; never as constant celebration"),
    ("Danger", DANGER, "Immediate physical or narrative threat"),
    ("Stealth", MOSS, "Reduced pace / concealment state"),
])
paragraph("", after=8)
heading("Accessibility rules", 2)
bullets([
    "Never communicate challenge outcome with color alone. Pair it with a word, icon change, haptic and audio response.",
    "Torch is designed for dark surfaces. On light surfaces, use Torch Dark for small text and thin strokes.",
    "Keep primary text at WCAG AA contrast or better; active-run information should target AAA where practical.",
    "Route colors require differentiated line patterns or terminal shapes for color-vision resilience.",
])

# 10 Grid
page_break()
title_block("09 / Layout", "Designed for a moving body.", "Worlds uses large targets, clear edges and predictable placement. Elegance comes from rhythm, not density.")
data_table(["Token", "Value", "Application"], [
    ("Base unit", "4 px", "All spacing resolves to the base unit"),
    ("Screen gutter", "20 px", "Standard phone content inset"),
    ("Section gap", "32 px", "Major content separation"),
    ("Card radius", "16 px", "World and recap cards"),
    ("Control radius", "999 px", "Primary run controls only"),
    ("Touch target", "48 px min", "Any interactive control"),
    ("Primary CTA", "56 px", "Start, resume and finish actions"),
    ("Bottom safety", "24 px + inset", "Thumb comfort and device safe area"),
], [1.45, 1.25, 4.45])
heading("Composition", 2)
bullets([
    "Use a 4-column mobile grid before the run and a simpler edge-anchored layout during the run.",
    "Allow one dominant object per screen: artwork, route, choice or recap statement.",
    "Avoid nested cards. A surface should not need a card inside a card to establish hierarchy.",
    "Place the most consequential action in the lower third, reachable by one thumb.",
    "During motion, no essential information may depend on horizontal scrolling, tiny map labels or long reading.",
])
callout("Glance test", "At jogging motion and arm's length, a runner must be able to identify the current objective and primary action in under one second.")

# 11 Architecture
page_break()
title_block("10 / Product architecture", "Four places. One continuous story.", "The navigation is deliberately smaller than a standard fitness or streaming app. Worlds should not reveal organizational complexity to the runner.")
data_table(["Destination", "Purpose", "Primary content"], [
    ("Home", "Return to the story", "Continue card, suggested worlds, recent run"),
    ("Worlds", "Browse intentionally", "Curated collections, genres, creator worlds"),
    ("Journey", "See what changed", "Campaign timeline, relationships, discoveries"),
    ("You", "Manage the runner", "Run history, goals, audio, safety, account"),
], [1.2, 2.05, 3.9])
heading("Persistent control", 2)
paragraph("When a run is active, navigation disappears. A compact live-run pill may appear only if the runner temporarily leaves the run screen. It shows elapsed time, one state label and a return action - never a miniature dashboard.")
heading("Home hierarchy", 2)
numbered([
    "Continue the current campaign. This is the largest element and contains one decisive action.",
    "Choose the shape of today's run: 15, 30 or 45 minutes, recovery, interval or open run.",
    "Discover one editorially selected world or creator collection.",
    "Review the last consequence in a single sentence.",
])

# 12 Home/library
page_break()
title_block("11 / Home & library", "Show worlds, not content inventory.", "Discovery should feel curated and cinematic. The interface never advertises volume with endless rows of near-identical tiles.")
heading("Continue card", 2)
data_table(["Layer", "Specification"], [
    ("Artwork", "Full-bleed cinematic key art; subject offset from copy and controls."),
    ("Eyebrow", "CURRENT JOURNEY / world or campaign name"),
    ("Title", "The immediate story objective, not an episode number."),
    ("State line", "One remembered consequence: 'Mara is injured. The northern gate is closed.'"),
    ("Action", "Continue - 30 min; duration is part of the decision."),
], [1.5, 5.65])
heading("Collection logic", 2)
bullets([
    "Worlds Originals - the quality benchmark and default first-run experience.",
    "Creator Worlds - lead with creator name and a clear authored-by signal.",
    "For tonight - mood and duration, not algorithmic personality claims.",
    "Your unfinished stories - progress expressed as narrative state, not percentages.",
])
heading("Tile rules", 2)
paragraph("Use a 4:5 poster ratio for worlds and a 16:9 landscape ratio for featured campaigns. Never bake essential labels into artwork. Display duration, audio readiness and campaign status outside the image with consistent UI typography.")
callout("Editorial discipline", "If every world is featured, none is. One hero, one secondary collection and one continuation state are enough for the hackathon home screen.")

# 13 Detail/pre-run
page_break()
title_block("12 / World detail & pre-run", "Build desire, then remove friction.", "The story page sells the emotional promise. The pre-run sheet converts that promise into a safe, ready session.")
heading("World detail order", 2)
numbered([
    "Key art and title treatment.",
    "One-sentence hook written in-world.",
    "A 20-30 second voice preview with captions.",
    "What movement controls: direction, pace, stops or a specific combination.",
    "Campaign status or new-story entry point.",
    "Creator, cast and authored-world credits.",
])
heading("Pre-run sheet", 2)
data_table(["Row", "Behavior"], [
    ("Run shape", "15 / 30 / 45 min or distance. Recommend one based on story state."),
    ("Intensity", "Recovery / steady / challenge. Explain in physical terms."),
    ("Audio", "Headphones connected, voice preview passed, volume safe."),
    ("Location", "GPS ready and route sensitivity clear."),
    ("Safety", "One concise reminder; never theatricalize real-world risk."),
    ("Primary action", "Start story. Full-width, 56 px, Torch on Ink."),
], [1.4, 5.75])
callout("Start-state rule", "The first authored line should begin within two seconds of starting. Do not place a branded countdown, tutorial carousel or loading monologue between the runner and the world.")

# 14 Active run
page_break()
title_block("13 / Active run", "The route is the stage.", "The active experience is a map, an objective and an audio presence. Everything else is subordinate.")
heading("Default hierarchy", 2)
data_table(["Zone", "Contains", "Rule"], [
    ("Top", "Elapsed time + one optional metric", "Never more than two numbers at once."),
    ("Center", "Route and runner position", "Low-contrast cartographic base; the directional runner arrow remains the only live overlay."),
    ("Lower third", "Current in-world objective", "One line; disappears when understood."),
    ("Bottom", "Audio state, pause, safety", "Large controls; protected from accidental taps."),
], [1.15, 2.45, 3.55])
heading("Route treatment", 2)
bullets([
    "Completed route: 40% neutral line. Current path: Torch, 5-6 px. Alternatives: thinner, lower contrast.",
    "For a minimal active-run map, a stylized raster base may stand on its own. Art-direct it into Ink and muted copper, remove provider branding and generated labels, and retain only the directional runner arrow as a native overlay.",
    "At a branch, simplify nearby map detail and increase path separation before showing any text.",
    "The runner marker points in the direction of travel; it should not resemble a generic location pin.",
    "Story locations appear only when relevant and use authored names: North Gate, Alder Wood, Bell Tower.",
])
heading("Metric restraint", 2)
paragraph("Default to elapsed time and current pace. Distance is available by tapping the metric group or through an audio check-in. Heart rate appears only when it changes the designed challenge or safety state.")
callout("Screen-off contract", "Every essential story decision and challenge can be understood through audio and movement alone. The screen confirms; it does not command.")

# 15 choices/challenges
page_break()
title_block("14 / Choices & challenges", "Movement is the control surface.", "The UI visualizes what the world has already made understandable through dialogue. It must never turn the moment into a multiple-choice form.")
heading("Direction choice", 2)
numbered([
    "Character names the physical possibilities in-world.",
    "Map subtly reveals the corresponding route branches.",
    "Runner changes direction; the selected route brightens immediately.",
    "A short haptic confirms recognition.",
    "Character acknowledges the choice within the next natural breath.",
])
heading("Pace challenge", 2)
data_table(["Phase", "Visual", "Audio / haptic"], [
    ("Set", "A thin target band appears around current pace.", "Character creates urgency; no synthetic beep."),
    ("Respond", "Runner marker and route pulse once as pace changes.", "One light haptic at threshold entry."),
    ("Hold", "Band contracts; no confetti or score accumulation.", "World sound carries the tension."),
    ("Resolve", "One-word state: ESCAPED / SEEN / INJURED.", "Character reaction and consequence line."),
], [1.05, 3.05, 3.05])
heading("Failure", 2)
paragraph("Failure is presented with the same visual dignity as success. No red full-screen punishment, buzzer or retry modal. The interface states what changed, then the story continues.")
callout("Example consequence", "THE GATE CLOSED / Mara took the lower path alone. This is a story state, not a fitness grade.")

# 16 Audio
page_break()
title_block("15 / Audio presence", "Voice is the interface.", "Audio state must be legible without resembling a music player. The runner is hearing a living scene, not consuming a track.")
data_table(["State", "Visual behavior"], [
    ("Character speaking", "Small route symbol pulses once on phrase onset; character name appears only if needed."),
    ("Narration", "No avatar or chat bubble. A quiet waveform line may indicate live playback."),
    ("Director reacting", "Never label AI or generation. Show 'The world is responding...' only if latency exceeds the designed threshold."),
    ("Music / ambience", "No persistent track title. Credits are available after the run."),
    ("Connection degraded", "Protect immersion with authored fallback audio, then show a plain recovery status."),
], [1.7, 5.45])
heading("Latency behavior", 2)
bullets([
    "Under 400 ms: immediate acknowledgement feels live.",
    "400-1,200 ms: bridge with breath, ambience or a character reaction sound.",
    "Over 1,200 ms: play an authored compatible line rather than expose silence or a spinner.",
    "Never repeat a detected choice while waiting for dialogue. One acknowledgement is enough.",
])
heading("Controls", 2)
paragraph("Pause, volume and captions remain available. Pausing opens a protected sheet with Resume run as the primary action and End run & view summary as a clear secondary action. The system back control opens this sheet rather than abandoning the session. Skip-forward is not a primary control because story timing and movement are connected. If a runner misses a line, offer a single 'Repeat last line' action and preserve state.")

# 17 Recap
page_break()
title_block("16 / Recap & progression", "Remember the run as a story.", "The recap is the payoff for the entire product promise. Lead with consequence; let fitness data support it.")
heading("Recap order", 2)
numbered([
    "A single authored headline based on the ending state: You reached the bell tower alone.",
    "A short spoken recap in the voices of characters who survived the run.",
    "A route map annotated with two or three meaningful story moments.",
    "Consequences: relationship, injury, resource, discovery and unresolved threat.",
    "Run facts: duration, distance and pace, visually quieter than the story result.",
    "One next-story hook with an honest minimum recommended run length.",
])
heading("Campaign timeline", 2)
data_table(["Entry type", "Example"], [
    ("Decision", "You entered Alder Wood instead of returning to the village."),
    ("Performance", "You outran the riders for 01:42."),
    ("Consequence", "Mara was injured at the northern gate."),
    ("Discovery", "The crest beneath the bridge matches the king's seal."),
    ("Relationship", "Ilan trusts you enough to reveal the hidden road."),
], [1.5, 5.65])
callout("Shareable artifact", "Generate a restrained story card: world art, one consequence line and a simplified route. Do not turn the recap into a generic fitness achievement template.")

page_break()
title_block("16A / Early-run summary", "Close the workout. Keep the story.", "Ending a physical run before a story beat is complete opens a functional summary, distinct from the completed-story recap. It may lead with distance, duration and average pace because the runner is closing a workout, but it must immediately confirm the saved story checkpoint and the next authored objective.")
heading("Summary order", 2)
numbered([
    "Run facts - Distance is dominant; duration and average pace are secondary. No calories, badges or score.",
    "Story state - Show a named checkpoint plus segmented moments reached; use a count, not a percentage or completion grade.",
    "Continuation - Resume story returns to run planning and promises continuation from the saved checkpoint.",
])
heading("State handoff", 2)
data_table(["Moment", "Required behavior"], [
    ("Pause", "Freeze workout telemetry and story audio. Resume run is primary; ending names the summary destination."),
    ("End", "Snapshot elapsed time, distance, average pace and the current authored story node."),
    ("Summary", "State that the run is saved without implying the current story beat was completed."),
    ("Resume", "Return to pre-run planning. Reset physical stats; continue the same saved story node and run state."),
], [1.35, 5.8])
callout("Continuation contract", "Ending a run closes the workout, not the campaign. The next run starts from zero physically and from the saved checkpoint narratively.")

# 18 Components
page_break()
title_block("17 / Component language", "Fewer components. Stronger states.", "A coherent MVP can be built from a small, rigorously specified set. Each component must have a reason to exist during motion.")
data_table(["Component", "Required states", "Notes"], [
    ("Primary button", "default / pressed / loading / disabled", "Torch fill; action verb first; no icon unless essential."),
    ("World card", "new / active / complete / locked", "Artwork-led; state outside image; one metadata line."),
    ("Route line", "past / current / alternate / selected / blocked", "Color plus line pattern or terminal shape."),
    ("Runner marker", "moving / stopped / uncertain GPS", "Direction-aware; uncertainty shown as radius, not alarm."),
    ("Objective label", "enter / active / resolved", "One sentence; auto-dismiss after comprehension window."),
    ("Audio presence", "speaking / ambience / buffering / fallback", "Never a chat bubble or assistant avatar."),
    ("Pause sheet", "paused / resuming / ending", "Resume is primary; ending names the summary destination."),
    ("Early-run summary", "saved / continuing", "Run facts first, then the named story checkpoint."),
    ("Consequence card", "positive / adverse / unresolved", "Same structure and dignity across outcomes."),
    ("Metric cluster", "default / challenge / safety", "Maximum one primary and two secondary values."),
], [1.35, 2.25, 3.55])
heading("Iconography", 2)
paragraph("Use a 2 px rounded-line icon set with optical rather than mathematical centering. Icons identify familiar actions only: pause, volume, captions, location and safety. Story concepts should be named, illustrated or voiced - not reduced to a generic icon library.")

# 19 Motion/haptics
page_break()
title_block("18 / Motion & haptics", "Movement should feel causal.", "Animation is used to connect the runner's action to the world's response. It should never create spectacle that competes with actual movement.")
data_table(["Event", "Motion", "Duration"], [
    ("Route selected", "Alternate route fades; selected line draws forward from runner.", "220-320 ms"),
    ("Pace threshold", "Target band settles around current pace; no continuous bounce.", "180 ms"),
    ("Objective enters", "Fade + 8 px upward translation.", "200 ms"),
    ("Consequence lands", "Short hold, then text reveal in two beats.", "400-600 ms"),
    ("Audio phrase onset", "One subtle pulse of the route symbol.", "160 ms"),
], [1.55, 4.25, 1.35])
heading("Haptic grammar", 2)
bullets([
    "Recognition: one light tap when movement is confidently understood.",
    "Success: two restrained taps, close together.",
    "Threat or safety: one firm tap; reserve it so it keeps meaning.",
    "Failure: no punitive buzz. Use the consequence's own audio and one grounded tap if needed.",
])
heading("Reduced motion", 2)
paragraph("All state changes remain understandable through contrast, shape, text, haptic and audio. Reduced-motion mode removes route drawing, parallax and pulsing while preserving instantaneous state recognition.")

# 20 imagery
page_break()
title_block("19 / Art direction", "Human stakes at landscape scale.", "World artwork establishes genre and emotional promise. It should look commissioned, cinematic and specific - never like a generic AI fantasy wallpaper.")
heading("Image principles", 2)
bullets([
    "Show a decisive place and a human point of view, not a collage of every character and threat.",
    "Use real directional composition: paths, light and gaze should lead toward the title or action area.",
    "Leave intentional negative space for responsive crops and interface copy.",
    "Favor tactile atmosphere - weather, distance, material, breath, light - over excessive detail.",
    "Maintain one art director, color script and character bible per world so assets feel authored across screens.",
])
heading("Crop system", 2)
data_table(["Ratio", "Use", "Composition requirement"], [
    ("4:5", "World poster", "Subject center-left or center-right; title-safe upper quarter."),
    ("16:9", "Featured campaign", "Narrative action on one side; copy-safe opposite third."),
    ("1:1", "Audio / share", "Readable silhouette and one focal relationship."),
    ("9:16", "Launch / transition", "Depth and forward direction; safe center corridor."),
], [1.0, 1.75, 4.4])
callout("AI imagery guardrail", "If generative tools are used during the hackathon, treat outputs as raw production material: art-direct, composite, retouch and enforce character continuity. Never ship an unedited first-generation image as key art.")

# 21 Prototype sequence
page_break()
title_block("20 / Hackathon sequence", "Design the proof, not the platform.", "The demo should make the interaction undeniable before explaining the technology. These screens and transitions deserve nearly all available polish.")
data_table(["Beat", "On screen", "What the audience understands"], [
    ("1 / Select", "One premium world and a 30-minute recommendation.", "This is curated entertainment."),
    ("2 / Begin", "Key art gives way to route; voice starts immediately.", "Audio is the primary interface."),
    ("3 / Choose", "Three routes emerge as the character describes them.", "Movement selects the path."),
    ("4 / React", "The chosen line brightens; dialogue names the choice.", "The world noticed instantly."),
    ("5 / Challenge", "Pace target appears inside the route state.", "Physical performance matters."),
    ("6 / Consequence", "Failure or success changes a character state.", "Failure creates story."),
    ("7 / Stop", "An optional discovery becomes available while stationary.", "Stopping is also input."),
    ("8 / Recap", "Route, choices and character consequence resolve together.", "The story remembers this exact run."),
], [1.15, 3.0, 3.0])
heading("Polish priority", 2)
paragraph("Spend disproportionate effort on audio transition timing, the route-choice acknowledgement, the pace-state change and the personalized recap. A shallow library can still feel intentional; a delayed reaction cannot.")

# 22 red flags/handoff
page_break()
title_block("21 / Design review", "Protect the idea from familiar shortcuts.", "Use this page as the final critique checklist before a screen, prototype or demo build is accepted.")
data_table(["Reject if…", "Because…"], [
    ("The active run resembles a dense fitness dashboard.", "The story has lost visual priority."),
    ("A choice appears as three tappable answer cards.", "Movement is no longer the interface."),
    ("The UI says AI, generated, event or branch.", "The directing system has broken immersion."),
    ("Every state uses Torch or a dramatic animation.", "The signal color and motion no longer carry meaning."),
    ("Failure is framed as a score loss or retry prompt.", "Consequence has become punishment."),
    ("World art determines button colors or core navigation.", "The platform loses learned consistency."),
    ("Copy sounds motivational, epic or generic.", "The authored voice has been replaced by product hype."),
    ("A completed-story recap leads with calories, pace or badges.", "The run did not become a story; stats-first treatment is reserved for an early-run summary."),
], [3.45, 3.7])
heading("MVP handoff tokens", 2)
paragraph("Ink #0B0C0E  /  Paper #F4F2ED  /  Torch #FF4F2E  /  Torch Dark #C92D14  /  Ash #686A70  /  Line #D8D6D0", size=9, color=INK, bold=True, after=7)
paragraph("Arial: Bold, Regular  /  Arial Narrow Bold: performance numerals only", size=9, color=INK, bold=True, after=7)
paragraph("Spacing: 4, 8, 12, 16, 20, 24, 32, 40, 48  /  Radii: 8, 16, 999  /  Touch: 48 minimum, 56 primary", size=9, color=INK, bold=True, after=12)
callout("Final test", "If the UI can be mistaken for an audiobook app, add agency. If it can be mistaken for a fantasy game, remove decoration. If it can be mistaken for a run tracker, restore the story.")

# 23 Sources
page_break()
title_block("Appendix", "Reference posture", "This direction borrows interaction discipline, not trade dress. Worlds should learn from proven run-time patterns while retaining its own typography, color, content model and narrative behavior.")
heading("Primary references", 2)
data_table(["Reference", "What to study", "Source"], [
    ("Nike Run Club", "Guided-run discovery, start flow, glanceable run metrics and confident restraint.", "nike.com/nrc-app"),
    ("Apple Workouts", "Active-workout hierarchy, safety and prioritization of essential metrics.", "developer.apple.com/design/human-interface-guidelines/workouts"),
    ("Zombies, Run!", "Screen-off narrative fitness behavior, mission framing and audio continuity.", "support.zombiesrungame.com"),
    ("Premium streaming", "Commissioned key art, continuation hierarchy and editorial curation.", "Pattern reference only; do not clone layout or branding."),
], [1.25, 3.35, 2.55])
heading("Originality boundary", 2)
paragraph("Do not reproduce another product's proprietary artwork, icons, exact screen compositions, copy, branded typefaces or distinctive trade dress. The useful lesson from NRC is disciplined hierarchy: one clear action, strong typography, restrained color and low cognitive load while moving.")
heading("Decision status", 2)
data_table(["Decision", "Status"], [
    ("Core visual direction", "Recommended for prototype"),
    ("Torch accent color", "Recommended; verify on target devices"),
    ("Arial", "Recommended for the hackathon working system"),
    ("Arial Narrow", "Hackathon numeric role; replace or license for distribution"),
    ("Wordmark customization", "Requires designer refinement before public launch"),
    ("Functional color accessibility", "Requires contrast audit in final components"),
], [2.75, 4.4])
paragraph("End of direction / Worlds v1.0", size=8.5, color=MUTED, bold=True, before=12, after=0)

doc.save(OUT)
print(OUT.resolve())
